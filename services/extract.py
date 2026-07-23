"""Question-bank text extraction (docx/pdf).

PDF extraction is copied from flashcard-generator/services/extract.py
(PyMuPDF text, with a page-image vision fallback for text-poor/scanned
pages) — generalized here from "one chapter" to "one question bank",
which is the same underlying operation.
"""

from pathlib import Path

import fitz
from docx import Document

TEXT_POOR_THRESHOLD = 120
MAX_VISION_PAGES = 24
VISION_DPI = 100
MAX_TEXT_CHARS = 200_000


def extract_pdf(pdf_path: str | Path) -> dict:
    """Return {mode, text, images, page_count} for a question-bank PDF."""
    doc = fitz.open(str(pdf_path))
    try:
        page_texts = [page.get_text("text").strip() for page in doc]
        total_chars = sum(len(t) for t in page_texts)
        avg_chars = total_chars / max(len(doc), 1)

        if avg_chars >= TEXT_POOR_THRESHOLD:
            text = "\n\n".join(
                f"[Page {i + 1}]\n{t}" for i, t in enumerate(page_texts) if t)
            return {"mode": "text", "text": text[:MAX_TEXT_CHARS],
                    "images": [], "page_count": len(doc)}

        images: list[bytes] = []
        for page in doc[:MAX_VISION_PAGES]:
            pix = page.get_pixmap(dpi=VISION_DPI)
            images.append(pix.tobytes("png"))
        return {"mode": "vision",
                "text": "\n\n".join(t for t in page_texts if t)[:MAX_TEXT_CHARS],
                "images": images, "page_count": len(doc)}
    finally:
        doc.close()


def extract_docx(docx_path: str | Path) -> dict:
    """Return {mode, text, images, page_count} for a question-bank docx.
    Tables are included (many question banks put options in a table)."""
    doc = Document(str(docx_path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return {"mode": "text", "text": text[:MAX_TEXT_CHARS], "images": [], "page_count": None}


def extract_document(path: str | Path) -> dict:
    path = Path(path)
    if path.suffix.lower() == ".pdf":
        return extract_pdf(path)
    if path.suffix.lower() in (".docx", ".doc"):
        return extract_docx(path)
    raise ValueError(f"unsupported question-bank file type: {path.suffix}")
